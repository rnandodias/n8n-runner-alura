"""
Geracao e manipulacao de documentos DOCX para revisao de artigos.
Inclui modelos Pydantic, helpers de formatacao e a funcao principal generate_docx.
"""
import re
import tempfile
from io import BytesIO
from typing import List, Optional
from urllib.parse import urljoin

import httpx
from fastapi import HTTPException
from fastapi.responses import Response
from PIL import Image as PILImage
from pydantic import BaseModel, field_validator

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================================
# MODELOS PYDANTIC
# ============================================================================

class TextSegment(BaseModel):
    text: str
    link: Optional[str] = None
    bold: Optional[bool] = False
    italic: Optional[bool] = False


class ContentItem(BaseModel):
    type: str
    level: Optional[int] = None
    text: Optional[str] = None
    segments: Optional[List[TextSegment]] = None
    ordered: Optional[bool] = False
    items: Optional[List] = None
    language: Optional[str] = None
    content: Optional[str] = None
    url: Optional[str] = None
    alt: Optional[str] = None
    width: Optional[int] = None
    height: Optional[int] = None
    headers: Optional[List[str]] = None
    rows: Optional[List[List[str]]] = None
    cite: Optional[str] = None

    @field_validator('segments', mode='before')
    @classmethod
    def filter_none_segments(cls, v):
        if v is None:
            return None
        return [seg for seg in v if seg is not None]

    @field_validator('items', mode='before')
    @classmethod
    def filter_none_items(cls, v):
        if v is None:
            return None
        return [item for item in v if item is not None]


class ArticleMetadata(BaseModel):
    title: Optional[str] = None
    author: Optional[str] = None
    publishDate: Optional[str] = None


class GenerateDocxPayload(BaseModel):
    metadata: ArticleMetadata
    content: List[ContentItem]
    filename: Optional[str] = "documento.docx"
    base_url: Optional[str] = None

    @field_validator('content', mode='before')
    @classmethod
    def filter_none_content(cls, v):
        if v is None:
            return []
        return [item for item in v if item is not None]


# ============================================================================
# HELPERS DE FORMATACAO DOCX
# ============================================================================

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0066CC')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')
    rPr.append(sz)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def convert_relative_url(url: str, base_url: str) -> str:
    if not url:
        return url
    if url.startswith('http://') or url.startswith('https://'):
        return url
    if not base_url:
        return url
    return urljoin(base_url, url)


def download_image(url: str) -> Optional[BytesIO]:
    try:
        with httpx.Client(timeout=30, follow_redirects=True) as client:
            response = client.get(url)
            response.raise_for_status()
            return BytesIO(response.content)
    except Exception as e:
        print(f"Erro ao baixar imagem {url}: {e}")
        return None


def convert_image_for_docx(image_bytes: Optional[BytesIO]) -> Optional[BytesIO]:
    """
    Converte imagem para formato compativel com python-docx.
    SVG -> PNG, WEBP animado -> GIF, outros -> PNG.
    """
    if image_bytes is None:
        return None

    SUPPORTED_FORMATS = {'PNG', 'JPEG', 'GIF', 'BMP', 'TIFF', 'JPG'}

    try:
        image_bytes.seek(0)
        raw = image_bytes.read(200)
        image_bytes.seek(0)
        if b'<svg' in raw or b'<?xml' in raw[:10]:
            try:
                import cairosvg
                image_bytes.seek(0)
                png_bytes = cairosvg.svg2png(file_obj=image_bytes)
                print(f"  [CONV] SVG -> PNG (cairosvg)")
                return BytesIO(png_bytes)
            except Exception as e:
                print(f"  [ERRO] Conversao SVG->PNG: {e}")
                return None
    except Exception:
        image_bytes.seek(0)

    try:
        image_bytes.seek(0)
        img = PILImage.open(image_bytes)

        if img.format and img.format.upper() in SUPPORTED_FORMATS:
            image_bytes.seek(0)
            return image_bytes

        if img.format == 'WEBP' and getattr(img, 'is_animated', False):
            print(f"  [CONV] WEBP animado -> GIF ({img.n_frames} frames)")
            return _convert_animated_webp_to_gif(img)

        print(f"  [CONV] {img.format or 'unknown'} -> PNG")
        return _convert_to_png(img)

    except Exception as e:
        print(f"  [ERRO] Conversao de imagem: {e}")
        return None


def _convert_animated_webp_to_gif(img) -> BytesIO:
    """Converte WEBP animado para GIF preservando animacao."""
    frames = []
    durations = []

    try:
        while True:
            frame = img.copy()

            if frame.mode in ('RGBA', 'LA'):
                background = PILImage.new('RGBA', frame.size, (255, 255, 255, 255))
                if frame.mode == 'RGBA':
                    background.paste(frame, mask=frame.split()[3])
                else:
                    background.paste(frame)
                frame = background.convert('RGB').convert('P', palette=PILImage.ADAPTIVE, colors=256)
            elif frame.mode != 'P':
                frame = frame.convert('P', palette=PILImage.ADAPTIVE, colors=256)

            frames.append(frame)
            durations.append(img.info.get('duration', 100))
            img.seek(img.tell() + 1)
    except EOFError:
        pass

    if not frames:
        raise ValueError("Nenhum frame extraido do WEBP animado")

    gif_buffer = BytesIO()
    frames[0].save(
        gif_buffer,
        format='GIF',
        save_all=True,
        append_images=frames[1:],
        duration=durations,
        loop=0
    )
    gif_buffer.seek(0)
    return gif_buffer


def _convert_to_png(img) -> BytesIO:
    """Converte imagem para PNG."""
    if img.mode in ('RGBA', 'LA', 'P'):
        if img.mode == 'P':
            img = img.convert('RGBA')
    elif img.mode not in ('RGB', 'L'):
        img = img.convert('RGB')

    png_buffer = BytesIO()
    img.save(png_buffer, format='PNG')
    png_buffer.seek(0)
    return png_buffer


def get_image_dimensions_from_bytes(image_bytes: BytesIO) -> tuple:
    try:
        image_bytes.seek(0)
        img = PILImage.open(image_bytes)
        width, height = img.size
        image_bytes.seek(0)
        return width, height
    except:
        return None, None


def set_paragraph_shading(paragraph, color: str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    paragraph._p.get_or_add_pPr().append(shading)


def add_left_border(paragraph, color: str = '0066CC', width: int = 24):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(width))
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def process_list_item_content_docx(doc, li, paragraph):
    """Processa conteúdo de item de lista no DOCX."""
    if li is None:
        return
    if isinstance(li, dict):
        if 'segments' in li and li['segments']:
            for seg in li['segments']:
                if seg is None or not isinstance(seg, dict):
                    continue
                seg_text = seg.get('text', '') or ''
                seg_link = seg.get('link')
                seg_bold = seg.get('bold', False)
                seg_italic = seg.get('italic', False)

                if seg_link:
                    add_hyperlink(paragraph, seg_text, seg_link)
                else:
                    run = paragraph.add_run(seg_text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    if seg_bold:
                        run.bold = True
                    if seg_italic:
                        run.italic = True
        elif 'text' in li and li['text']:
            run = paragraph.add_run(str(li['text']))
            run.font.name = 'Arial'
            run.font.size = Pt(12)
    elif li:
        run = paragraph.add_run(str(li))
        run.font.name = 'Arial'
        run.font.size = Pt(12)


def process_nested_list_docx(doc, items, ordered=False, indent_level=0):
    """Processa lista aninhada no DOCX."""
    if not items:
        return
    markers = ["• ", "◦ ", "▪ ", "- "]

    for idx, li in enumerate(items):
        if li is None:
            continue
        list_para = doc.add_paragraph()
        list_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if ordered:
            prefix = f"{idx + 1}. "
        else:
            prefix = markers[min(indent_level, len(markers) - 1)]

        prefix_run = list_para.add_run(prefix)
        prefix_run.font.name = 'Arial'
        prefix_run.font.size = Pt(12)

        process_list_item_content_docx(doc, li, list_para)

        base_indent = 0.5
        list_para.paragraph_format.left_indent = Inches(base_indent + (indent_level * 0.3))
        list_para.space_after = Pt(3)

        if isinstance(li, dict) and 'sublist' in li and li['sublist']:
            sublist = li['sublist']
            process_nested_list_docx(
                doc,
                sublist.get('items', []),
                sublist.get('ordered', False),
                indent_level + 1
            )


# ============================================================================
# GERACAO DE DOCX
# ============================================================================

async def generate_docx(payload: GenerateDocxPayload):
    """Gera documento Word (.docx) a partir de JSON estruturado."""
    try:
        print(f"Gerando DOCX: {payload.metadata.title or 'Sem titulo'}")

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)

        if payload.metadata.title:
            title_para = doc.add_heading(payload.metadata.title, level=1)
            for run in title_para.runs:
                run.bold = True
                run.font.size = Pt(28)
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(33, 37, 41)
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_para.space_after = Pt(6)

        meta_parts = []
        if payload.metadata.author:
            meta_parts.append(f"Por {payload.metadata.author}")
        if payload.metadata.publishDate:
            meta_parts.append(payload.metadata.publishDate)

        if meta_parts:
            meta_para = doc.add_paragraph()
            meta_run = meta_para.add_run(" • ".join(meta_parts))
            meta_run.italic = True
            meta_run.font.size = Pt(11)
            meta_run.font.color.rgb = RGBColor(102, 102, 102)
            meta_para.space_after = Pt(12)

        doc.add_paragraph("_" * 80)

        for item in payload.content:
            if item is None:
                continue

            if item.type == "heading" and item.text:
                spacer = doc.add_paragraph()
                spacer.space_after = Pt(0)
                spacer.space_before = Pt(6)

                level = item.level if item.level else 2
                heading_para = doc.add_heading(item.text, level=level)

                for run in heading_para.runs:
                    run.bold = True
                    run.font.name = 'Arial'

                    if level == 2:
                        run.font.size = Pt(16)
                        run.font.color.rgb = RGBColor(33, 37, 41)
                    elif level == 3:
                        run.font.size = Pt(14)
                        run.font.color.rgb = RGBColor(44, 62, 80)
                    elif level == 4:
                        run.font.size = Pt(13)
                        run.font.color.rgb = RGBColor(52, 73, 94)
                    elif level == 5:
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(60, 80, 100)
                    else:
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(70, 90, 110)

                heading_para.space_before = Pt(12)
                heading_para.space_after = Pt(6)

            elif item.type == "paragraph":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                if item.segments:
                    for seg in item.segments:
                        if seg is None:
                            continue
                        if seg.link:
                            add_hyperlink(para, seg.text or '', seg.link)
                        else:
                            run = para.add_run(seg.text or '')
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                            if seg.bold:
                                run.bold = True
                            if seg.italic:
                                run.italic = True
                elif item.text:
                    run = para.add_run(item.text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

                para.space_after = Pt(6)

            elif item.type == "list" and item.items:
                process_nested_list_docx(doc, item.items, item.ordered or False, indent_level=0)
                doc.add_paragraph()

            elif item.type == "blockquote":
                if item.segments:
                    quote_para = doc.add_paragraph()
                    quote_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    for seg in item.segments:
                        if seg is None:
                            continue
                        if seg.link:
                            add_hyperlink(quote_para, seg.text or '', seg.link)
                        else:
                            run = quote_para.add_run(seg.text or '')
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                            run.italic = True
                            run.font.color.rgb = RGBColor(85, 85, 85)

                    add_left_border(quote_para, color='0066CC', width=24)
                    quote_para.paragraph_format.left_indent = Inches(0.3)
                    quote_para.space_before = Pt(6)
                    quote_para.space_after = Pt(6)

                elif item.text:
                    quote_para = doc.add_paragraph()
                    quote_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    run = quote_para.add_run(item.text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    run.italic = True
                    run.font.color.rgb = RGBColor(85, 85, 85)

                    add_left_border(quote_para, color='0066CC', width=24)
                    quote_para.paragraph_format.left_indent = Inches(0.3)
                    quote_para.space_before = Pt(6)
                    quote_para.space_after = Pt(6)

                if item.cite:
                    cite_para = doc.add_paragraph()
                    cite_run = cite_para.add_run(f"— {item.cite}")
                    cite_run.font.name = 'Arial'
                    cite_run.font.size = Pt(10)
                    cite_run.italic = True
                    cite_run.font.color.rgb = RGBColor(120, 120, 120)
                    cite_para.paragraph_format.left_indent = Inches(0.5)
                    cite_para.space_after = Pt(12)

            elif item.type == "code" and item.content:
                if item.language:
                    lang_para = doc.add_paragraph()
                    lang_run = lang_para.add_run(f" {item.language.upper()} ")
                    lang_run.font.name = 'Consolas'
                    lang_run.font.size = Pt(9)
                    lang_run.font.color.rgb = RGBColor(255, 255, 255)
                    set_paragraph_shading(lang_para, '2d2d2d')
                    lang_para.space_after = Pt(0)

                for line in item.content.split('\n'):
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run(line if line else ' ')
                    code_run.font.name = 'Consolas'
                    code_run.font.size = Pt(10)
                    code_run.font.color.rgb = RGBColor(51, 51, 51)
                    set_paragraph_shading(code_para, 'F8F8F8')
                    code_para.paragraph_format.left_indent = Inches(0.2)
                    code_para.space_after = Pt(0)
                    code_para.space_before = Pt(0)

                doc.add_paragraph().space_after = Pt(12)

            elif item.type == "image" and item.url:
                image_url = convert_relative_url(item.url, payload.base_url)
                print(f"Baixando imagem: {image_url[:80]}...")
                image_data = download_image(image_url)

                if image_data:
                    image_data = convert_image_for_docx(image_data)

                if image_data:
                    try:
                        orig_width, orig_height = get_image_dimensions_from_bytes(image_data)
                        max_width_cm = 15

                        if orig_width and orig_height:
                            width_cm = orig_width / 96 * 2.54
                            height_cm = orig_height / 96 * 2.54
                            if width_cm > max_width_cm:
                                ratio = max_width_cm / width_cm
                                width_cm = max_width_cm
                                height_cm = height_cm * ratio
                            img_para = doc.add_paragraph()
                            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = img_para.add_run()
                            run.add_picture(image_data, width=Cm(width_cm))
                        else:
                            img_para = doc.add_paragraph()
                            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = img_para.add_run()
                            run.add_picture(image_data, width=Cm(max_width_cm))

                        img_para.space_after = Pt(6)

                        if item.alt and len(item.alt) > 5:
                            caption_para = doc.add_paragraph()
                            caption_run = caption_para.add_run(item.alt)
                            caption_run.italic = True
                            caption_run.font.size = Pt(10)
                            caption_run.font.color.rgb = RGBColor(102, 102, 102)
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_para.space_after = Pt(12)

                        print(f"Imagem adicionada")
                    except Exception as img_error:
                        print(f"Erro ao processar imagem: {img_error}")

            elif item.type == "table" and item.headers and item.rows:
                print(f"Adicionando tabela com {len(item.rows)} linhas...")

                num_cols = len(item.headers)
                num_rows = len(item.rows) + 1

                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'

                header_row = table.rows[0]
                for idx, header_text in enumerate(item.headers):
                    cell = header_row.cells[idx]
                    cell.text = header_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'E0E0E0')
                    cell._tc.get_or_add_tcPr().append(shading)

                for row_idx, row_data in enumerate(item.rows):
                    row = table.rows[row_idx + 1]
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = row.cells[col_idx]
                            cell.text = str(cell_text) if cell_text else ""
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(10)

                doc.add_paragraph().space_after = Pt(12)

        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        filename = payload.filename
        if not filename.endswith('.docx'):
            filename += '.docx'
        filename = re.sub(r'[^a-zA-Z0-9\s\-_.]', '', filename)

        print(f"DOCX gerado: {filename}")

        return Response(
            content=doc_buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    except Exception as e:
        print(f"Erro ao gerar DOCX: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erro ao gerar DOCX: {str(e)}")
