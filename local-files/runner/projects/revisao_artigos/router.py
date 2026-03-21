"""
Router FastAPI para o projeto de Revisao Automatica de Artigos.
Prefix: /revisao/artigos
"""
import json
import os
import tempfile
from typing import Optional

import httpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel

from core.llm_client import criar_cliente_llm
from core.track_changes import aplicar_comentarios_docx, aplicar_revisoes_docx
from projects.revisao_artigos.docx_builder import (
    ArticleMetadata,
    ContentItem,
    GenerateDocxPayload,
    generate_docx,
)
from projects.revisao_artigos.prompts import (
    formatar_prompt_imagem,
    formatar_prompt_seo,
    formatar_prompt_tecnico,
    formatar_prompt_texto,
)
from projects.revisao_artigos.scraping import extract_article_content

router = APIRouter(prefix="/revisao/artigos")


# ============================================================================
# MODELOS PYDANTIC
# ============================================================================

class ExtractArticlePayload(BaseModel):
    url: str


class RevisaoItem(BaseModel):
    tipo: str
    acao: str
    texto_original: str
    texto_novo: Optional[str] = ""
    justificativa: str = ""


class AplicarRevisoesPayload(BaseModel):
    docx_url: Optional[str] = None
    docx_base64: Optional[str] = None
    revisoes: list
    autor: str = "Agente IA Revisor"


class ExtrairTextoDocxPayload(BaseModel):
    docx_url: Optional[str] = None
    docx_base64: Optional[str] = None


class RevisaoAgentPayload(BaseModel):
    docx_url: Optional[str] = None
    docx_base64: Optional[str] = None
    provider: str = "anthropic"
    guia_seo_url: Optional[str] = None
    url_artigo: Optional[str] = ""
    titulo: Optional[str] = ""
    data_publicacao: Optional[str] = ""


class RevisaoImagemPayload(BaseModel):
    docx_url: Optional[str] = None
    docx_base64: Optional[str] = None
    provider: str = "anthropic"
    url_artigo: str
    titulo: Optional[str] = ""


# ============================================================================
# HELPERS
# ============================================================================

async def obter_docx_bytes(docx_url: Optional[str], docx_base64: Optional[str], http_client=None) -> bytes:
    """Obtém bytes do DOCX a partir de URL ou base64."""
    if docx_base64:
        import base64
        if docx_base64.startswith("data:"):
            docx_base64 = docx_base64.split(",", 1)[1]
        return base64.b64decode(docx_base64)
    elif docx_url:
        if http_client is None:
            async with httpx.AsyncClient(timeout=60) as client:
                resp = await client.get(docx_url)
                resp.raise_for_status()
                return resp.content
        else:
            resp = await http_client.get(docx_url)
            resp.raise_for_status()
            return resp.content
    else:
        raise ValueError("Deve fornecer docx_url ou docx_base64")


def _is_image_caption(para) -> bool:
    """Detecta se um paragrafo e uma legenda de imagem."""
    if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return False

    runs = para.runs
    if not runs:
        return False

    for run in runs:
        if not run.italic:
            return False
        if run.font.size and run.font.size != Pt(10):
            return False
        if run.font.color.rgb:
            if run.font.color.rgb != RGBColor(102, 102, 102):
                return False

    return True


def _extrair_texto_para_revisao(docx_path: str, incluir_legendas: bool = False) -> tuple:
    """Extrai texto estruturado de um DOCX para analise. Retorna (texto_completo, titulo)."""
    doc = Document(docx_path)
    paragrafos = []
    texto_parts = []
    titulo = ""
    idx = 0

    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue

        if _is_image_caption(para):
            if incluir_legendas:
                paragrafos.append({
                    "indice": idx,
                    "texto": texto,
                    "tipo": "image_caption"
                })
                texto_parts.append(f"[P{idx}|IMAGE_CAPTION] {texto}")
                idx += 1
            continue

        estilo = para.style.name if para.style else "Normal"
        tipo = "paragraph"
        if "heading" in estilo.lower():
            tipo = estilo.lower().replace(" ", "")

        if not titulo and "heading1" in tipo:
            titulo = texto

        paragrafos.append({
            "indice": idx,
            "texto": texto,
            "tipo": tipo,
            "tamanho": len(texto)
        })

        texto_parts.append(f"[P{idx}|{tipo.upper()}] {texto}")
        idx += 1

    texto_completo = "\n\n".join(texto_parts)
    return texto_completo, titulo


# ============================================================================
# ENDPOINTS - HTML PARA DOCX
# ============================================================================

@router.post("/html-to-docx")
async def html_to_docx(payload: ExtractArticlePayload):
    """Pipeline completo: extrai artigo de URL e gera DOCX em uma única chamada."""
    try:
        print(f"Pipeline HTML -> DOCX: {payload.url}")

        with httpx.Client(timeout=30, follow_redirects=True) as client:
            response = client.get(payload.url)
            response.raise_for_status()
            html = response.text

        article_data = extract_article_content(html, payload.url)
        print(f"Extraido: {article_data['stats']}")

        docx_payload = GenerateDocxPayload(
            metadata=ArticleMetadata(**article_data['metadata']),
            content=[ContentItem(**item) for item in article_data['content']],
            filename=article_data['filename'],
            base_url=article_data['base_url']
        )

        return await generate_docx(docx_payload)

    except Exception as e:
        print(f"Erro no pipeline: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erro no pipeline: {str(e)}")


# ============================================================================
# ENDPOINTS - REVISAO
# ============================================================================

@router.post("/extrair-texto")
async def revisao_extrair_texto(payload: ExtrairTextoDocxPayload):
    """Extrai texto estruturado de um documento DOCX para analise."""
    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)
        paragrafos = []
        texto_parts = []
        idx = 0

        for para in doc.paragraphs:
            texto = para.text.strip()
            if not texto:
                continue

            estilo = para.style.name if para.style else "Normal"
            tipo = "paragraph"
            if "heading" in estilo.lower():
                tipo = estilo.lower().replace(" ", "")

            paragrafos.append({
                "indice": idx,
                "texto": texto,
                "tipo": tipo,
                "tamanho": len(texto)
            })

            texto_parts.append(f"[P{idx}|{tipo.upper()}] {texto}")
            idx += 1

        return {
            "paragrafos": paragrafos,
            "texto_completo": "\n\n".join(texto_parts),
            "total_paragrafos": idx
        }
    finally:
        os.unlink(tmp_path)


@router.post("/aplicar")
async def revisao_aplicar(payload: AplicarRevisoesPayload):
    """Aplica revisoes a um documento DOCX com Track Changes."""
    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        input_path = tmp.name

    output_path = input_path.replace(".docx", "_revisado.docx")

    try:
        revisoes_list = [r.model_dump() if hasattr(r, 'model_dump') else r for r in payload.revisoes]

        resultado = aplicar_revisoes_docx(
            input_path,
            output_path,
            revisoes_list,
            payload.autor
        )

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="documento_revisado.docx",
            headers={
                "X-Total-Revisoes": str(resultado["total_revisoes"]),
                "X-Aplicadas": str(resultado["aplicadas"]),
                "X-Falhas": str(resultado["falhas"]),
                "X-Comentarios": str(resultado["comentarios"])
            }
        )
    except Exception as e:
        raise HTTPException(500, f"Erro ao aplicar revisoes: {str(e)}")
    finally:
        if os.path.exists(input_path):
            os.unlink(input_path)


@router.post("/aplicar-json")
async def revisao_aplicar_json(
    docx_url: str = Form(...),
    revisoes: str = Form(...),
    autor: str = Form("Agente IA Revisor")
):
    """Aplica revisoes via Form (compativel com n8n HTTP Request)."""
    async with httpx.AsyncClient(timeout=60.0) as client:
        resp = await client.get(docx_url)
        if resp.status_code != 200:
            raise HTTPException(400, f"Erro ao baixar documento: {resp.status_code}")

    try:
        revisoes_list = json.loads(revisoes)
    except json.JSONDecodeError as e:
        raise HTTPException(400, f"JSON invalido: {e}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(resp.content)
        input_path = tmp.name

    output_path = input_path.replace(".docx", "_revisado.docx")

    try:
        resultado = aplicar_revisoes_docx(
            input_path,
            output_path,
            revisoes_list,
            autor
        )

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="documento_revisado.docx",
            headers={
                "X-Total-Revisoes": str(resultado["total_revisoes"]),
                "X-Aplicadas": str(resultado["aplicadas"]),
                "X-Falhas": str(resultado["falhas"]),
                "X-Comentarios": str(resultado["comentarios"])
            }
        )
    except Exception as e:
        raise HTTPException(500, f"Erro ao aplicar revisoes: {str(e)}")
    finally:
        if os.path.exists(input_path):
            os.unlink(input_path)


@router.post("/agente-seo")
async def revisao_agente_seo(payload: RevisaoAgentPayload):
    """Executa o agente de revisao SEO."""
    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    guia_seo = "Use boas praticas gerais de SEO para conteudo tecnico educacional."
    if payload.guia_seo_url:
        try:
            async with httpx.AsyncClient(timeout=60.0) as http_client:
                guia_resp = await http_client.get(payload.guia_seo_url)
                if guia_resp.status_code == 200:
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                        tmp.write(guia_resp.content)
                        guia_path = tmp.name
                    guia_doc = Document(guia_path)
                    guia_seo = "\n".join([p.text for p in guia_doc.paragraphs if p.text.strip()])
                    os.unlink(guia_path)
        except Exception as e:
            print(f"Aviso: Nao foi possivel carregar guia SEO: {e}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo = payload.titulo or titulo_extraido

        system_prompt, user_prompt, artigo_context = formatar_prompt_seo(
            conteudo=conteudo,
            titulo=titulo,
            url=payload.url_artigo or "",
            guia_seo=guia_seo
        )

        llm_client = criar_cliente_llm(provider=payload.provider)
        resposta = llm_client.gerar_resposta(system_prompt, user_prompt, artigo_context=artigo_context)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "SEO"

        return {
            "tipo": "SEO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente SEO: {str(e)}")
    finally:
        os.unlink(tmp_path)


@router.post("/agente-tecnico")
async def revisao_agente_tecnico(payload: RevisaoAgentPayload):
    """Executa o agente de revisao TECNICA."""
    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo = payload.titulo or titulo_extraido

        system_prompt, user_prompt, artigo_context = formatar_prompt_tecnico(
            conteudo=conteudo,
            titulo=titulo,
            url=payload.url_artigo or "",
            data_publicacao=payload.data_publicacao or ""
        )

        llm_client = criar_cliente_llm(provider=payload.provider)
        resposta = llm_client.gerar_resposta_com_busca(system_prompt, user_prompt, artigo_context=artigo_context)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "TECNICO"

        return {
            "tipo": "TECNICO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente TECNICO: {str(e)}")
    finally:
        os.unlink(tmp_path)


@router.post("/agente-texto")
async def revisao_agente_texto(payload: RevisaoAgentPayload):
    """Executa o agente de revisao TEXTUAL."""
    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo = payload.titulo or titulo_extraido

        system_prompt, user_prompt, artigo_context = formatar_prompt_texto(
            conteudo=conteudo,
            titulo=titulo,
            url=payload.url_artigo or ""
        )

        llm_client = criar_cliente_llm(provider=payload.provider)
        resposta = llm_client.gerar_resposta(system_prompt, user_prompt, artigo_context=artigo_context)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "TEXTO"

        return {
            "tipo": "TEXTO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente TEXTO: {str(e)}")
    finally:
        os.unlink(tmp_path)


@router.post("/agente-seo-form")
async def revisao_agente_seo_form(
    file: UploadFile = File(...),
    provider: str = Form("anthropic"),
    url_artigo: str = Form(""),
    titulo: str = Form(""),
    guia_seo_file: Optional[UploadFile] = File(None),
    palavras_chave: str = Form("")
):
    """Executa o agente de revisao SEO via upload de arquivo."""
    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    guia_path = None

    try:
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo_final = titulo or titulo_extraido

        guia_seo = "Use boas praticas gerais de SEO para conteudo tecnico educacional."
        if guia_seo_file and guia_seo_file.filename:
            guia_bytes = await guia_seo_file.read()
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_guia:
                tmp_guia.write(guia_bytes)
                tmp_guia.flush()
                guia_path = tmp_guia.name

            guia_doc = Document(guia_path)
            guia_seo = "\n".join([p.text for p in guia_doc.paragraphs if p.text.strip()])

        palavras_formatadas = "Nenhuma palavra-chave especifica fornecida. Use seu conhecimento de SEO."
        if palavras_chave and palavras_chave.strip():
            keywords = [kw.strip() for kw in palavras_chave.replace('\n', ',').split(',') if kw.strip()]
            if keywords:
                palavras_formatadas = "\n".join([f"- {kw}" for kw in keywords])

        system_prompt, user_prompt, artigo_context = formatar_prompt_seo(
            conteudo=conteudo,
            titulo=titulo_final,
            url=url_artigo,
            guia_seo=guia_seo,
            palavras_chave=palavras_formatadas
        )

        llm_client = criar_cliente_llm(provider=provider)
        resposta = llm_client.gerar_resposta(system_prompt, user_prompt, artigo_context=artigo_context)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "SEO"

        return {
            "tipo": "SEO",
            "total_sugestoes": len(revisoes),
            "palavras_chave_usadas": palavras_chave.strip() if palavras_chave else None,
            "guia_seo_fornecido": bool(guia_seo_file and guia_seo_file.filename),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente SEO: {str(e)}")
    finally:
        os.unlink(tmp_path)
        if guia_path and os.path.exists(guia_path):
            os.unlink(guia_path)


@router.post("/agente-tecnico-form")
async def revisao_agente_tecnico_form(
    file: UploadFile = File(...),
    provider: str = Form("anthropic"),
    url_artigo: str = Form(""),
    titulo: str = Form(""),
    data_publicacao: str = Form("")
):
    """Executa o agente de revisao TECNICA via upload de arquivo."""
    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo_final = titulo or titulo_extraido

        system_prompt, user_prompt, artigo_context = formatar_prompt_tecnico(
            conteudo=conteudo,
            titulo=titulo_final,
            url=url_artigo,
            data_publicacao=data_publicacao
        )

        llm_client = criar_cliente_llm(provider=provider)
        resposta = llm_client.gerar_resposta_com_busca(system_prompt, user_prompt, artigo_context=artigo_context)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "TECNICO"

        return {
            "tipo": "TECNICO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente TECNICO: {str(e)}")
    finally:
        os.unlink(tmp_path)


@router.post("/agente-texto-form")
async def revisao_agente_texto_form(
    file: UploadFile = File(...),
    provider: str = Form("anthropic"),
    url_artigo: str = Form(""),
    titulo: str = Form("")
):
    """Executa o agente de revisao TEXTUAL via upload de arquivo."""
    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo_final = titulo or titulo_extraido

        system_prompt, user_prompt, artigo_context = formatar_prompt_texto(
            conteudo=conteudo,
            titulo=titulo_final,
            url=url_artigo
        )

        llm_client = criar_cliente_llm(provider=provider)
        resposta = llm_client.gerar_resposta(system_prompt, user_prompt, artigo_context=artigo_context)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "TEXTO"

        return {
            "tipo": "TEXTO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente TEXTO: {str(e)}")
    finally:
        os.unlink(tmp_path)


@router.post("/agente-imagem")
async def revisao_agente_imagem(payload: RevisaoImagemPayload):
    """Executa o agente de revisao de IMAGENS (visao multimodal)."""
    from datetime import datetime

    if not payload.url_artigo:
        raise HTTPException(400, "url_artigo e obrigatorio para o agente de imagem")

    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(docx_bytes)
            tmp.flush()
            tmp_path = tmp.name

        try:
            conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
            titulo_final = payload.titulo or titulo_extraido

            print(f"Extraindo imagens de: {payload.url_artigo}")
            with httpx.Client(timeout=30, follow_redirects=True) as http_client:
                response = http_client.get(payload.url_artigo)
                response.raise_for_status()
                html = response.text

            article_data = extract_article_content(html, payload.url_artigo)

            imagens = [item for item in article_data.get('content', []) if item.get('type') == 'image']
            print(f"{len(imagens)} imagens encontradas")

            if not imagens:
                return {
                    "tipo": "IMAGEM",
                    "total_sugestoes": 0,
                    "revisoes": [],
                    "mensagem": "Nenhuma imagem encontrada no artigo"
                }

            data_atual = datetime.now().strftime("%d/%m/%Y")
            system_prompt, user_prompt, artigo_context = formatar_prompt_imagem(
                conteudo=conteudo,
                imagens=imagens,
                titulo=titulo_final,
                url=payload.url_artigo,
                data_atual=data_atual
            )

            llm_client = criar_cliente_llm(provider=payload.provider)
            resposta = llm_client.gerar_resposta_com_imagens_e_busca(
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                imagens=imagens,
                artigo_context=artigo_context
            )

            revisoes = llm_client.extrair_json(resposta)
            revisoes_validas = []
            for rev in revisoes:
                if isinstance(rev, dict):
                    rev["tipo"] = "IMAGEM"
                    revisoes_validas.append(rev)

            return {
                "tipo": "IMAGEM",
                "total_sugestoes": len(revisoes_validas),
                "total_imagens": len(imagens),
                "revisoes": revisoes_validas
            }

        finally:
            os.unlink(tmp_path)

    except httpx.HTTPStatusError as e:
        status = e.response.status_code
        if status >= 500:
            raise HTTPException(502, f"Erro no servidor do artigo ({status}): {e.request.url}")
        else:
            raise HTTPException(400, f"Erro ao buscar URL do artigo ({status}): {e.request.url}")
    except httpx.HTTPError as e:
        raise HTTPException(502, f"Erro de conexao ao buscar artigo: {str(e)}")
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Erro no agente IMAGEM: {str(e)}")


@router.post("/agente-imagem-form")
async def revisao_agente_imagem_form(
    file: UploadFile = File(...),
    url_artigo: str = Form(...),
    provider: str = Form("anthropic"),
    titulo: str = Form("")
):
    """Executa o agente de revisao de IMAGENS via upload de arquivo."""
    from datetime import datetime

    if not url_artigo:
        raise HTTPException(400, "url_artigo e obrigatorio para o agente de imagem")

    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo_final = titulo or titulo_extraido

        print(f"Extraindo imagens de: {url_artigo}")
        with httpx.Client(timeout=30, follow_redirects=True) as http_client:
            response = http_client.get(url_artigo)
            response.raise_for_status()
            html = response.text

        article_data = extract_article_content(html, url_artigo)

        imagens = [item for item in article_data.get('content', []) if item.get('type') == 'image']
        print(f"{len(imagens)} imagens encontradas")

        if not imagens:
            return {
                "tipo": "IMAGEM",
                "total_sugestoes": 0,
                "revisoes": [],
                "mensagem": "Nenhuma imagem encontrada no artigo"
            }

        data_atual = datetime.now().strftime("%d/%m/%Y")
        system_prompt, user_prompt, artigo_context = formatar_prompt_imagem(
            conteudo=conteudo,
            imagens=imagens,
            titulo=titulo_final,
            url=url_artigo,
            data_atual=data_atual
        )

        llm_client = criar_cliente_llm(provider=provider)
        resposta = llm_client.gerar_resposta_com_imagens_e_busca(
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            imagens=imagens,
            artigo_context=artigo_context
        )

        revisoes = llm_client.extrair_json(resposta)
        revisoes_validas = []
        for rev in revisoes:
            if isinstance(rev, dict):
                rev["tipo"] = "IMAGEM"
                revisoes_validas.append(rev)

        return {
            "tipo": "IMAGEM",
            "total_sugestoes": len(revisoes_validas),
            "total_imagens": len(imagens),
            "revisoes": revisoes_validas
        }

    except httpx.HTTPStatusError as e:
        status = e.response.status_code
        if status >= 500:
            raise HTTPException(502, f"Erro no servidor do artigo ({status}): {e.request.url}")
        else:
            raise HTTPException(400, f"Erro ao buscar URL do artigo ({status}): {e.request.url}")
    except httpx.HTTPError as e:
        raise HTTPException(502, f"Erro de conexao ao buscar artigo: {str(e)}")
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Erro no agente IMAGEM: {str(e)}")
    finally:
        os.unlink(tmp_path)


@router.post("/aplicar-form")
async def revisao_aplicar_form(
    file: UploadFile = File(...),
    revisoes: str = Form(...),
    autor: str = Form("Agente IA Revisor")
):
    """Aplica revisoes a um documento DOCX via upload de arquivo."""
    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        input_path = tmp.name

    output_path = input_path.replace(".docx", "_revisado.docx")

    try:
        revisoes_list = json.loads(revisoes)

        resultado = aplicar_revisoes_docx(
            input_path,
            output_path,
            revisoes_list,
            autor
        )

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="documento_revisado.docx",
            headers={
                "X-Total-Revisoes": str(resultado["total_revisoes"]),
                "X-Aplicadas": str(resultado["aplicadas"]),
                "X-Falhas": str(resultado["falhas"]),
                "X-Comentarios": str(resultado["comentarios"])
            }
        )
    except json.JSONDecodeError as e:
        raise HTTPException(400, f"JSON de revisoes invalido: {str(e)}")
    except Exception as e:
        raise HTTPException(500, f"Erro ao aplicar revisoes: {str(e)}")
    finally:
        if os.path.exists(input_path):
            os.unlink(input_path)


@router.post("/aplicar-comentarios-form")
async def revisao_aplicar_comentarios_form(
    file: UploadFile = File(...),
    revisoes: str = Form(...),
    autor: str = Form("Agente IA Revisor")
):
    """Aplica SOMENTE comentarios a um documento DOCX via upload de arquivo."""
    original_filename = file.filename or "documento.docx"

    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        input_path = tmp.name

    output_path = input_path.replace(".docx", "_comentado.docx")

    try:
        revisoes_list = json.loads(revisoes)

        resultado = aplicar_comentarios_docx(
            input_path,
            output_path,
            revisoes_list,
            autor
        )

        stats = resultado.get('estatisticas', {})

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=original_filename,
            headers={
                "X-Total-Comentarios": str(resultado.get('total_comentarios', 0)),
                "X-Match-Exato": str(stats.get('exato', 0)),
                "X-Match-Normalizado": str(stats.get('normalizado', 0)),
                "X-Match-Fuzzy": str(stats.get('fuzzy', 0)),
                "X-Match-Paragrafo": str(stats.get('paragrafo', 0)),
                "X-Match-Falhas": str(stats.get('falha', 0)),
            }
        )
    except json.JSONDecodeError as e:
        raise HTTPException(400, f"JSON de revisoes invalido: {str(e)}")
    except Exception as e:
        raise HTTPException(500, f"Erro ao aplicar comentarios: {str(e)}")
    finally:
        if os.path.exists(input_path):
            os.unlink(input_path)
