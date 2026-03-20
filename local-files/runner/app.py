# ============================================================================
# RUNNER ALURA - FastAPI - Agente de Revisao de Artigos
# ============================================================================

import os
import json
import re
import tempfile
from pathlib import Path
from typing import List, Optional
from io import BytesIO

# Third-party
from unidecode import unidecode
from bs4 import BeautifulSoup, NavigableString
from urllib.parse import urljoin
import httpx

# FastAPI
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.responses import Response, FileResponse
from pydantic import BaseModel, field_validator

# DOCX
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# PIL - DEVE ser importado ANTES do UNO para evitar conflito de imports
from PIL import Image as PILImage

# Track Changes OOXML
from track_changes import aplicar_revisoes_docx, aplicar_comentarios_docx

# LLM e Prompts para agentes de revisao
from llm_client import criar_cliente_llm
from prompts_revisao import (
    formatar_prompt_seo,
    formatar_prompt_tecnico,
    formatar_prompt_texto,
    formatar_prompt_imagem
)


# ============================================================================
# MODELOS PYDANTIC
# ============================================================================

# --- DOCX ---
class TextSegment(BaseModel):
    text: str
    link: Optional[str] = None
    bold: Optional[bool] = False
    italic: Optional[bool] = False


class ContentItem(BaseModel):
    type: str
    level: Optional[int] = None
    text: Optional[str] = None
    segments: Optional[List[TextSegment]] = None
    ordered: Optional[bool] = False
    items: Optional[List] = None
    language: Optional[str] = None
    content: Optional[str] = None
    url: Optional[str] = None
    alt: Optional[str] = None
    width: Optional[int] = None
    height: Optional[int] = None
    headers: Optional[List[str]] = None
    rows: Optional[List[List[str]]] = None
    cite: Optional[str] = None

    @field_validator('segments', mode='before')
    @classmethod
    def filter_none_segments(cls, v):
        if v is None:
            return None
        return [seg for seg in v if seg is not None]

    @field_validator('items', mode='before')
    @classmethod
    def filter_none_items(cls, v):
        if v is None:
            return None
        return [item for item in v if item is not None]


class ArticleMetadata(BaseModel):
    title: Optional[str] = None
    author: Optional[str] = None
    publishDate: Optional[str] = None


class GenerateDocxPayload(BaseModel):
    metadata: ArticleMetadata
    content: List[ContentItem]
    filename: Optional[str] = "documento.docx"
    base_url: Optional[str] = None

    @field_validator('content', mode='before')
    @classmethod
    def filter_none_content(cls, v):
        if v is None:
            return []
        return [item for item in v if item is not None]


class ExtractArticlePayload(BaseModel):
    url: str


# --- Revisao com Track Changes (OOXML) ---
class RevisaoItem(BaseModel):
    tipo: str  # "SEO", "TECNICO", "TEXTO", "IMAGEM"
    acao: str  # "substituir", "deletar", "inserir", "comentario"
    texto_original: str
    texto_novo: Optional[str] = ""
    justificativa: str = ""


class AplicarRevisoesPayload(BaseModel):
    docx_url: Optional[str] = None
    docx_base64: Optional[str] = None
    revisoes: List[RevisaoItem]
    autor: str = "Agente IA Revisor"


class ExtrairTextoDocxPayload(BaseModel):
    docx_url: Optional[str] = None
    docx_base64: Optional[str] = None


class RevisaoAgentPayload(BaseModel):
    docx_url: Optional[str] = None
    docx_base64: Optional[str] = None
    provider: str = "anthropic"
    guia_seo_url: Optional[str] = None
    url_artigo: Optional[str] = ""
    titulo: Optional[str] = ""
    data_publicacao: Optional[str] = ""


class RevisaoImagemPayload(BaseModel):
    docx_url: Optional[str] = None
    docx_base64: Optional[str] = None
    provider: str = "anthropic"
    url_artigo: str
    titulo: Optional[str] = ""


# ============================================================================
# HELPERS - GERAIS
# ============================================================================

async def obter_docx_bytes(docx_url: Optional[str], docx_base64: Optional[str], http_client=None) -> bytes:
    """Obtém bytes do DOCX a partir de URL ou base64."""
    if docx_base64:
        import base64
        if docx_base64.startswith("data:"):
            docx_base64 = docx_base64.split(",", 1)[1]
        return base64.b64decode(docx_base64)
    elif docx_url:
        if http_client is None:
            async with httpx.AsyncClient(timeout=60) as client:
                resp = await client.get(docx_url)
                resp.raise_for_status()
                return resp.content
        else:
            resp = await http_client.get(docx_url)
            resp.raise_for_status()
            return resp.content
    else:
        raise ValueError("Deve fornecer docx_url ou docx_base64")


# ============================================================================
# HELPERS - EXTRACAO DE ARTIGOS (BeautifulSoup)
# ============================================================================

def is_banner_or_promotional(element):
    """Verifica se elemento é banner/propaganda."""
    parent_a = element.find_parent('a') if element.name != 'a' else element
    if parent_a and parent_a.get('href'):
        href = parent_a.get('href', '')
        promo_patterns = [
            '/escola-', '/formacao-', '/planos-', '/curso-online',
            '/empresas', 'cursos.alura.com.br/loginForm',
            'utm_source=blog', 'utm_medium=banner', 'utm_campaign=',
            '/carreiras/', '/pos-tech'
        ]
        for pattern in promo_patterns:
            if pattern in href:
                return True

    if element.name == 'img':
        src = element.get('src', '').lower()
        alt = element.get('alt', '').lower()
        if any(x in src for x in ['matricula-escola', 'saiba-mais', 'banner']):
            return True
        if 'banner' in alt:
            return True

    return False


def is_site_chrome(element):
    """Verifica se elemento faz parte do chrome do site."""
    if element.find_parent(['nav', 'footer', 'aside']):
        return True

    parent_header = element.find_parent('header')
    if parent_header:
        if parent_header.find('a', href=lambda x: x and '/carreiras' in x):
            return True

    if element.find_parent(class_=lambda x: x and 'cosmos-author' in str(x)):
        return True

    if element.find_parent(class_=lambda x: x and 'social-media' in str(x)):
        return True
    if element.find_parent(class_=lambda x: x and 'cosmos-container-social' in str(x)):
        return True

    if element.name == 'p':
        text = element.get_text(strip=True).lower()
        if text == 'compartilhe':
            return True

    return False


def is_decorative_element(element):
    """Verifica se é elemento decorativo."""
    if element.name == 'img':
        src = element.get('src', '').lower()
        classes = element.get('class', [])

        if 'cosmos-image' in classes:
            return False

        if 'cdn-wcsm.alura.com.br' in src:
            return False

        decorative_patterns = [
            '/assets/img/header/', '/assets/img/home/', '/assets/img/caelum',
            '/assets/img/footer/', '/assets/img/ecossistema/',
            'arrow-', 'return-', 'icon', 'avatar',
            'gravatar.com/avatar', 'gnarususercontent.com.br'
        ]

        for pattern in decorative_patterns:
            if pattern in src:
                return True

        if '.svg' in src and '/assets/' in src:
            return True

        width = element.get('width')
        if width:
            try:
                if int(width) < 50:
                    return True
            except ValueError:
                pass

    return False


def get_text_preserving_spaces(element):
    """Extrai texto preservando espaços entre elementos inline."""
    texts = []
    for child in element.descendants:
        if isinstance(child, NavigableString):
            texts.append(str(child))
    result = ''.join(texts)
    result = re.sub(r'\s+', ' ', result)
    return result.strip()


def extract_text_with_formatting(element, base_url):
    """Extrai texto preservando formatação (links, bold, italic)."""
    segments = []

    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip():
                segments.append({"text": text})

        elif child.name == 'a':
            href = child.get('href', '')
            text = child.get_text()
            if text.strip():
                if href and not href.startswith('http') and not href.startswith('#'):
                    href = urljoin(base_url, href)
                segments.append({"text": text, "link": href if href else None})

        elif child.name in ['strong', 'b']:
            inner_a = child.find('a')
            inner_em = child.find(['em', 'i'])

            if inner_a:
                href = inner_a.get('href', '')
                if href and not href.startswith('http') and not href.startswith('#'):
                    href = urljoin(base_url, href)
                text = child.get_text()
                if text.strip():
                    segments.append({"text": text, "link": href, "bold": True})
            elif inner_em:
                for subchild in child.children:
                    if isinstance(subchild, NavigableString):
                        text = str(subchild)
                        if text.strip():
                            segments.append({"text": text, "bold": True})
                    elif subchild.name in ['em', 'i']:
                        em_a = subchild.find('a')
                        if em_a:
                            href = em_a.get('href', '')
                            if href and not href.startswith('http') and not href.startswith('#'):
                                href = urljoin(base_url, href)
                            segments.append({"text": subchild.get_text(), "link": href, "bold": True, "italic": True})
                        else:
                            segments.append({"text": subchild.get_text(), "bold": True, "italic": True})
                    elif subchild.name == 'a':
                        href = subchild.get('href', '')
                        if href and not href.startswith('http') and not href.startswith('#'):
                            href = urljoin(base_url, href)
                        segments.append({"text": subchild.get_text(), "link": href, "bold": True})
            else:
                text = child.get_text()
                if text.strip():
                    segments.append({"text": text, "bold": True})

        elif child.name in ['em', 'i']:
            inner_a = child.find('a')
            if inner_a:
                href = inner_a.get('href', '')
                if href and not href.startswith('http') and not href.startswith('#'):
                    href = urljoin(base_url, href)
                text = child.get_text()
                if text.strip():
                    segments.append({"text": text, "link": href, "italic": True})
            else:
                text = child.get_text()
                if text.strip():
                    segments.append({"text": text, "italic": True})

        elif child.name == 'code':
            text = child.get_text()
            if text.strip():
                segments.append({"text": f"`{text}`", "bold": True})

        elif child.name == 'p':
            inner_segments = extract_text_with_formatting(child, base_url)
            segments.extend(inner_segments)

        elif child.name in ['span', 'mark', 'u']:
            inner_segments = extract_text_with_formatting(child, base_url)
            segments.extend(inner_segments)

        elif child.name == 'br':
            segments.append({"text": "\n"})

        elif child.name in ['sup', 'sub']:
            text = child.get_text()
            if text.strip():
                segments.append({"text": text})

        else:
            text = child.get_text()
            if text.strip():
                segments.append({"text": text})

    return segments


def process_list_items(ul_or_ol, base_url, ordered=False):
    """Processa itens de lista, incluindo listas aninhadas."""
    items = []

    for li in ul_or_ol.find_all('li', recursive=False):
        item = {}
        sublist = li.find(['ul', 'ol'], recursive=False)

        if sublist:
            sublist_copy = sublist.extract()
            segments = extract_text_with_formatting(li, base_url)
            li.append(sublist_copy)

            if segments:
                has_formatting = any(
                    seg.get('link') or seg.get('bold') or seg.get('italic')
                    for seg in segments
                )

                if has_formatting:
                    item['segments'] = segments
                elif len(segments) == 1:
                    item['text'] = segments[0].get('text', '').strip()
                else:
                    item['text'] = ''.join(seg.get('text', '') for seg in segments).strip()

            sub_ordered = sublist_copy.name == 'ol'
            sub_items = process_list_items(sublist_copy, base_url, sub_ordered)
            if sub_items:
                item['sublist'] = {
                    'ordered': sub_ordered,
                    'items': sub_items
                }
        else:
            segments = extract_text_with_formatting(li, base_url)
            if segments:
                has_formatting = any(
                    seg.get('link') or seg.get('bold') or seg.get('italic')
                    for seg in segments
                )

                if has_formatting:
                    item['segments'] = segments
                elif len(segments) == 1:
                    item['text'] = segments[0].get('text', '').strip()
                else:
                    item['text'] = ''.join(seg.get('text', '') for seg in segments).strip()

        if item:
            items.append(item)

    return items


def extract_table(table_tag):
    """Extrai dados de tabela HTML."""
    headers = []
    rows = []

    thead = table_tag.find('thead')
    if thead:
        header_row = thead.find('tr')
        if header_row:
            headers = [th.get_text(strip=True) for th in header_row.find_all(['th', 'td'])]

    if not headers:
        first_row = table_tag.find('tr')
        if first_row:
            ths = first_row.find_all('th')
            if ths:
                headers = [th.get_text(strip=True) for th in ths]

    tbody = table_tag.find('tbody') or table_tag
    for tr in tbody.find_all('tr'):
        if tr.find('th') and not rows and headers:
            continue

        cells = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
        if cells and any(c for c in cells):
            rows.append(cells)

    return headers, rows


def extract_article_content(html: str, base_url: str) -> dict:
    """
    Extrai conteúdo estruturado de artigo Alura usando BeautifulSoup.
    100% determinístico, sem IA.
    """
    soup = BeautifulSoup(html, 'html.parser')

    for tag in soup.find_all(['script', 'style', 'noscript', 'svg', 'iframe']):
        tag.decompose()

    metadata = {
        'title': None,
        'author': None,
        'publishDate': None
    }
    content = []
    processed_elements = set()

    h1 = soup.find('h1')
    if h1:
        metadata['title'] = h1.get_text(strip=True)
        processed_elements.add(id(h1))

    date_pattern = re.compile(r'\d{2}/\d{2}/\d{4}')
    page_text = soup.get_text()
    date_match = date_pattern.search(page_text)
    if date_match:
        metadata['publishDate'] = date_match.group()

    author_candidates = []
    for img in soup.find_all('img'):
        src = img.get('src', '')
        alt = img.get('alt', '')
        if 'gravatar.com' in src or 'gnarususercontent.com.br' in src:
            if alt and len(alt) > 2 and not any(x in alt.lower() for x in ['logo', 'banner', 'alura']):
                author_candidates.append(alt)

    if author_candidates:
        metadata['author'] = author_candidates[0]

    main_content = soup.find('body') or soup
    stop_processing = False

    list_item_texts = set()
    for li in main_content.find_all('li'):
        li_text = li.get_text(strip=True)
        if li_text and len(li_text) > 10:
            list_item_texts.add(li_text)

    for element in main_content.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'p', 'ul', 'ol',
                                           'blockquote', 'pre', 'table', 'img', 'figure']):
        elem_id = id(element)
        if elem_id in processed_elements:
            continue
        processed_elements.add(elem_id)

        if is_site_chrome(element):
            continue
        if is_banner_or_promotional(element):
            continue
        if is_decorative_element(element):
            continue

        if element.name in ['h2', 'h3']:
            text = element.get_text(strip=True).lower()
            if any(x in text for x in ['leia também', 'artigos relacionados', 'veja outros artigos']):
                stop_processing = True

        if stop_processing:
            continue

        if element.name == 'h1':
            continue

        if element.name in ['h2', 'h3', 'h4', 'h5']:
            text = get_text_preserving_spaces(element)
            if text and len(text) > 1:
                if element.find_parent(class_=lambda x: x and 'toc' in x.lower() if x else False):
                    continue
                level = int(element.name[1])
                content.append({
                    'type': 'heading',
                    'level': level,
                    'text': text
                })

        elif element.name == 'p':
            text = element.get_text(strip=True)
            if not text:
                continue
            if text in list_item_texts:
                continue

            segments = extract_text_with_formatting(element, base_url)
            if segments:
                has_formatting = any(
                    seg.get('link') or seg.get('bold') or seg.get('italic')
                    for seg in segments
                )

                if not has_formatting and len(segments) == 1:
                    content.append({
                        'type': 'paragraph',
                        'text': segments[0].get('text', '').strip()
                    })
                else:
                    content.append({
                        'type': 'paragraph',
                        'segments': segments
                    })

        elif element.name in ['ul', 'ol']:
            if element.find_parent(['ul', 'ol']):
                continue

            ordered = element.name == 'ol'
            items = process_list_items(element, base_url, ordered)

            if items:
                content.append({
                    'type': 'list',
                    'ordered': ordered,
                    'items': items
                })

        elif element.name == 'blockquote':
            segments = extract_text_with_formatting(element, base_url)
            cite_tag = element.find('cite')
            cite = cite_tag.get_text(strip=True) if cite_tag else None

            if segments:
                blockquote_item = {'type': 'blockquote', 'segments': segments}
                if cite:
                    blockquote_item['cite'] = cite
                content.append(blockquote_item)

        elif element.name == 'pre':
            code_tag = element.find('code')
            if code_tag:
                code_content = code_tag.get_text()
                classes = code_tag.get('class', [])
                language = None
                for cls in classes:
                    if isinstance(cls, str):
                        if cls.startswith('language-'):
                            language = cls.replace('language-', '')
                            break
                        elif cls in ['python', 'javascript', 'java', 'sql', 'bash',
                                    'html', 'css', 'json', 'typescript', 'jsx', 'ruby',
                                    'go', 'rust', 'php', 'csharp', 'kotlin', 'swift']:
                            language = cls
                            break

                content.append({
                    'type': 'code',
                    'language': language,
                    'content': code_content
                })
            else:
                content.append({
                    'type': 'code',
                    'content': element.get_text()
                })

        elif element.name == 'table':
            headers, rows = extract_table(element)
            if headers or rows:
                content.append({
                    'type': 'table',
                    'headers': headers,
                    'rows': rows
                })

        elif element.name == 'img':
            src = element.get('src', '')
            if not src:
                continue
            if is_banner_or_promotional(element):
                continue
            if is_decorative_element(element):
                continue

            if not src.startswith('http'):
                src = urljoin(base_url, src)

            alt = element.get('alt', '')
            width = element.get('width')
            height = element.get('height')

            img_item = {
                'type': 'image',
                'url': src,
                'alt': alt
            }

            if width:
                try:
                    img_item['width'] = int(width)
                except:
                    pass
            if height:
                try:
                    img_item['height'] = int(height)
                except:
                    pass

            content.append(img_item)

        elif element.name == 'figure':
            img = element.find('img')
            if img:
                src = img.get('src', '')
                if not src:
                    continue

                if not src.startswith('http'):
                    src = urljoin(base_url, src)

                figcaption = element.find('figcaption')
                alt = figcaption.get_text(strip=True) if figcaption else img.get('alt', '')

                content.append({
                    'type': 'image',
                    'url': src,
                    'alt': alt
                })
                processed_elements.add(id(img))

    content = [item for item in content if item]

    stats = {}
    for item in content:
        item_type = item.get('type', 'unknown')
        stats[item_type] = stats.get(item_type, 0) + 1

    filename = metadata.get('title', 'documento') or 'documento'
    filename = unidecode(filename)
    filename = re.sub(r'[^a-zA-Z0-9\s-]', '', filename)
    filename = re.sub(r'\s+', '-', filename).strip('-')
    filename = filename[:80]
    filename = f"{filename}.docx"

    return {
        'metadata': metadata,
        'content': content,
        'filename': filename,
        'base_url': base_url,
        'stats': stats
    }


# ============================================================================
# HELPERS - GERACAO DE DOCX
# ============================================================================

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0066CC')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')
    rPr.append(sz)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def convert_relative_url(url: str, base_url: str) -> str:
    if not url:
        return url
    if url.startswith('http://') or url.startswith('https://'):
        return url
    if not base_url:
        return url
    return urljoin(base_url, url)


def download_image(url: str) -> Optional[BytesIO]:
    try:
        with httpx.Client(timeout=30, follow_redirects=True) as client:
            response = client.get(url)
            response.raise_for_status()
            return BytesIO(response.content)
    except Exception as e:
        print(f"Erro ao baixar imagem {url}: {e}")
        return None


def convert_image_for_docx(image_bytes: Optional[BytesIO]) -> Optional[BytesIO]:
    """
    Converte imagem para formato compativel com python-docx.

    - SVG -> PNG (via cairosvg)
    - WEBP animado -> GIF (preserva animacao)
    - WEBP estatico -> PNG
    - Outros formatos nao suportados -> PNG
    - Formatos suportados (PNG, JPEG, GIF, BMP, TIFF) -> retorna original
    """
    if image_bytes is None:
        return None

    SUPPORTED_FORMATS = {'PNG', 'JPEG', 'GIF', 'BMP', 'TIFF', 'JPG'}

    try:
        image_bytes.seek(0)
        raw = image_bytes.read(200)
        image_bytes.seek(0)
        if b'<svg' in raw or b'<?xml' in raw[:10]:
            try:
                import cairosvg
                image_bytes.seek(0)
                png_bytes = cairosvg.svg2png(file_obj=image_bytes)
                print(f"  [CONV] SVG -> PNG (cairosvg)")
                return BytesIO(png_bytes)
            except Exception as e:
                print(f"  [ERRO] Conversao SVG->PNG: {e}")
                return None
    except Exception:
        image_bytes.seek(0)

    try:
        image_bytes.seek(0)
        img = PILImage.open(image_bytes)

        if img.format and img.format.upper() in SUPPORTED_FORMATS:
            image_bytes.seek(0)
            return image_bytes

        if img.format == 'WEBP' and getattr(img, 'is_animated', False):
            print(f"  [CONV] WEBP animado -> GIF ({img.n_frames} frames)")
            return _convert_animated_webp_to_gif(img)

        print(f"  [CONV] {img.format or 'unknown'} -> PNG")
        return _convert_to_png(img)

    except Exception as e:
        print(f"  [ERRO] Conversao de imagem: {e}")
        return None


def _convert_animated_webp_to_gif(img) -> BytesIO:
    """Converte WEBP animado para GIF preservando animacao."""
    frames = []
    durations = []

    try:
        while True:
            frame = img.copy()

            if frame.mode in ('RGBA', 'LA'):
                background = PILImage.new('RGBA', frame.size, (255, 255, 255, 255))
                if frame.mode == 'RGBA':
                    background.paste(frame, mask=frame.split()[3])
                else:
                    background.paste(frame)
                frame = background.convert('RGB').convert('P', palette=PILImage.ADAPTIVE, colors=256)
            elif frame.mode != 'P':
                frame = frame.convert('P', palette=PILImage.ADAPTIVE, colors=256)

            frames.append(frame)
            durations.append(img.info.get('duration', 100))
            img.seek(img.tell() + 1)
    except EOFError:
        pass

    if not frames:
        raise ValueError("Nenhum frame extraido do WEBP animado")

    gif_buffer = BytesIO()
    frames[0].save(
        gif_buffer,
        format='GIF',
        save_all=True,
        append_images=frames[1:],
        duration=durations,
        loop=0
    )
    gif_buffer.seek(0)
    return gif_buffer


def _convert_to_png(img) -> BytesIO:
    """Converte imagem para PNG."""
    if img.mode in ('RGBA', 'LA', 'P'):
        if img.mode == 'P':
            img = img.convert('RGBA')
    elif img.mode not in ('RGB', 'L'):
        img = img.convert('RGB')

    png_buffer = BytesIO()
    img.save(png_buffer, format='PNG')
    png_buffer.seek(0)
    return png_buffer


def get_image_dimensions_from_bytes(image_bytes: BytesIO) -> tuple:
    try:
        image_bytes.seek(0)
        img = PILImage.open(image_bytes)
        width, height = img.size
        image_bytes.seek(0)
        return width, height
    except:
        return None, None


def set_paragraph_shading(paragraph, color: str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    paragraph._p.get_or_add_pPr().append(shading)


def add_left_border(paragraph, color: str = '0066CC', width: int = 24):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(width))
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def process_list_item_content_docx(doc, li, paragraph):
    """Processa conteúdo de item de lista no DOCX."""
    if li is None:
        return
    if isinstance(li, dict):
        if 'segments' in li and li['segments']:
            for seg in li['segments']:
                if seg is None or not isinstance(seg, dict):
                    continue
                seg_text = seg.get('text', '') or ''
                seg_link = seg.get('link')
                seg_bold = seg.get('bold', False)
                seg_italic = seg.get('italic', False)

                if seg_link:
                    add_hyperlink(paragraph, seg_text, seg_link)
                else:
                    run = paragraph.add_run(seg_text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    if seg_bold:
                        run.bold = True
                    if seg_italic:
                        run.italic = True
        elif 'text' in li and li['text']:
            run = paragraph.add_run(str(li['text']))
            run.font.name = 'Arial'
            run.font.size = Pt(12)
    elif li:
        run = paragraph.add_run(str(li))
        run.font.name = 'Arial'
        run.font.size = Pt(12)


def process_nested_list_docx(doc, items, ordered=False, indent_level=0):
    """Processa lista aninhada no DOCX."""
    if not items:
        return
    markers = ["• ", "◦ ", "▪ ", "- "]

    for idx, li in enumerate(items):
        if li is None:
            continue
        list_para = doc.add_paragraph()
        list_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if ordered:
            prefix = f"{idx + 1}. "
        else:
            prefix = markers[min(indent_level, len(markers) - 1)]

        prefix_run = list_para.add_run(prefix)
        prefix_run.font.name = 'Arial'
        prefix_run.font.size = Pt(12)

        process_list_item_content_docx(doc, li, list_para)

        base_indent = 0.5
        list_para.paragraph_format.left_indent = Inches(base_indent + (indent_level * 0.3))
        list_para.space_after = Pt(3)

        if isinstance(li, dict) and 'sublist' in li and li['sublist']:
            sublist = li['sublist']
            process_nested_list_docx(
                doc,
                sublist.get('items', []),
                sublist.get('ordered', False),
                indent_level + 1
            )


# ============================================================================
# APP
# ============================================================================

app = FastAPI()


# ============================================================================
# ENDPOINTS - GERAL
# ============================================================================

@app.get("/ping")
def ping():
    return {"ok": True, "service": "runner"}


# ============================================================================
# ENDPOINTS - HTML PARA DOCX
# ============================================================================

@app.post("/html-to-docx")
async def html_to_docx(payload: ExtractArticlePayload):
    """Pipeline completo: extrai artigo de URL e gera DOCX em uma única chamada."""
    try:
        print(f"Pipeline HTML -> DOCX: {payload.url}")

        with httpx.Client(timeout=30, follow_redirects=True) as client:
            response = client.get(payload.url)
            response.raise_for_status()
            html = response.text

        article_data = extract_article_content(html, payload.url)
        print(f"Extraido: {article_data['stats']}")

        docx_payload = GenerateDocxPayload(
            metadata=ArticleMetadata(**article_data['metadata']),
            content=[ContentItem(**item) for item in article_data['content']],
            filename=article_data['filename'],
            base_url=article_data['base_url']
        )

        return await generate_docx(docx_payload)

    except Exception as e:
        print(f"Erro no pipeline: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erro no pipeline: {str(e)}")


async def generate_docx(payload: GenerateDocxPayload):
    """Gera documento Word (.docx) a partir de JSON estruturado."""
    try:
        print(f"Gerando DOCX: {payload.metadata.title or 'Sem titulo'}")

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)

        if payload.metadata.title:
            title_para = doc.add_heading(payload.metadata.title, level=1)
            for run in title_para.runs:
                run.bold = True
                run.font.size = Pt(28)
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(33, 37, 41)
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_para.space_after = Pt(6)

        meta_parts = []
        if payload.metadata.author:
            meta_parts.append(f"Por {payload.metadata.author}")
        if payload.metadata.publishDate:
            meta_parts.append(payload.metadata.publishDate)

        if meta_parts:
            meta_para = doc.add_paragraph()
            meta_run = meta_para.add_run(" • ".join(meta_parts))
            meta_run.italic = True
            meta_run.font.size = Pt(11)
            meta_run.font.color.rgb = RGBColor(102, 102, 102)
            meta_para.space_after = Pt(12)

        doc.add_paragraph("_" * 80)

        for item in payload.content:
            if item is None:
                continue

            if item.type == "heading" and item.text:
                spacer = doc.add_paragraph()
                spacer.space_after = Pt(0)
                spacer.space_before = Pt(6)

                level = item.level if item.level else 2
                heading_para = doc.add_heading(item.text, level=level)

                for run in heading_para.runs:
                    run.bold = True
                    run.font.name = 'Arial'

                    if level == 2:
                        run.font.size = Pt(16)
                        run.font.color.rgb = RGBColor(44, 62, 80)
                    elif level == 3:
                        run.font.size = Pt(14)
                        run.font.color.rgb = RGBColor(52, 73, 94)
                    elif level == 4:
                        run.font.size = Pt(13)
                        run.font.color.rgb = RGBColor(60, 80, 100)
                    else:
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(70, 90, 110)

                heading_para.space_before = Pt(12)
                heading_para.space_after = Pt(6)

            elif item.type == "paragraph":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                if item.segments:
                    for seg in item.segments:
                        if seg is None:
                            continue
                        if seg.link:
                            add_hyperlink(para, seg.text or '', seg.link)
                        else:
                            run = para.add_run(seg.text or '')
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                            if seg.bold:
                                run.bold = True
                            if seg.italic:
                                run.italic = True
                elif item.text:
                    run = para.add_run(item.text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

                para.space_after = Pt(6)

            elif item.type == "list" and item.items:
                process_nested_list_docx(doc, item.items, item.ordered or False, indent_level=0)
                doc.add_paragraph()

            elif item.type == "blockquote":
                if item.segments:
                    quote_para = doc.add_paragraph()
                    quote_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    for seg in item.segments:
                        if seg is None:
                            continue
                        if seg.link:
                            add_hyperlink(quote_para, seg.text or '', seg.link)
                        else:
                            run = quote_para.add_run(seg.text or '')
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                            run.italic = True
                            run.font.color.rgb = RGBColor(85, 85, 85)

                    add_left_border(quote_para, color='0066CC', width=24)
                    quote_para.paragraph_format.left_indent = Inches(0.3)
                    quote_para.space_before = Pt(6)
                    quote_para.space_after = Pt(6)

                elif item.text:
                    quote_para = doc.add_paragraph()
                    quote_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    run = quote_para.add_run(item.text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    run.italic = True
                    run.font.color.rgb = RGBColor(85, 85, 85)

                    add_left_border(quote_para, color='0066CC', width=24)
                    quote_para.paragraph_format.left_indent = Inches(0.3)
                    quote_para.space_before = Pt(6)
                    quote_para.space_after = Pt(6)

                if item.cite:
                    cite_para = doc.add_paragraph()
                    cite_run = cite_para.add_run(f"— {item.cite}")
                    cite_run.font.name = 'Arial'
                    cite_run.font.size = Pt(10)
                    cite_run.italic = True
                    cite_run.font.color.rgb = RGBColor(120, 120, 120)
                    cite_para.paragraph_format.left_indent = Inches(0.5)
                    cite_para.space_after = Pt(12)

            elif item.type == "code" and item.content:
                if item.language:
                    lang_para = doc.add_paragraph()
                    lang_run = lang_para.add_run(f" {item.language.upper()} ")
                    lang_run.font.name = 'Consolas'
                    lang_run.font.size = Pt(9)
                    lang_run.font.color.rgb = RGBColor(255, 255, 255)
                    set_paragraph_shading(lang_para, '2d2d2d')
                    lang_para.space_after = Pt(0)

                for line in item.content.split('\n'):
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run(line if line else ' ')
                    code_run.font.name = 'Consolas'
                    code_run.font.size = Pt(10)
                    code_run.font.color.rgb = RGBColor(51, 51, 51)
                    set_paragraph_shading(code_para, 'F8F8F8')
                    code_para.paragraph_format.left_indent = Inches(0.2)
                    code_para.space_after = Pt(0)
                    code_para.space_before = Pt(0)

                doc.add_paragraph().space_after = Pt(12)

            elif item.type == "image" and item.url:
                image_url = convert_relative_url(item.url, payload.base_url)
                print(f"Baixando imagem: {image_url[:80]}...")
                image_data = download_image(image_url)

                if image_data:
                    image_data = convert_image_for_docx(image_data)

                if image_data:
                    try:
                        orig_width, orig_height = get_image_dimensions_from_bytes(image_data)
                        max_width_cm = 15

                        if orig_width and orig_height:
                            width_cm = orig_width / 96 * 2.54
                            height_cm = orig_height / 96 * 2.54
                            if width_cm > max_width_cm:
                                ratio = max_width_cm / width_cm
                                width_cm = max_width_cm
                                height_cm = height_cm * ratio
                            img_para = doc.add_paragraph()
                            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = img_para.add_run()
                            run.add_picture(image_data, width=Cm(width_cm))
                        else:
                            img_para = doc.add_paragraph()
                            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = img_para.add_run()
                            run.add_picture(image_data, width=Cm(max_width_cm))

                        img_para.space_after = Pt(6)

                        if item.alt and len(item.alt) > 5:
                            caption_para = doc.add_paragraph()
                            caption_run = caption_para.add_run(item.alt)
                            caption_run.italic = True
                            caption_run.font.size = Pt(10)
                            caption_run.font.color.rgb = RGBColor(102, 102, 102)
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_para.space_after = Pt(12)

                        print(f"Imagem adicionada")
                    except Exception as img_error:
                        print(f"Erro ao processar imagem: {img_error}")

            elif item.type == "table" and item.headers and item.rows:
                print(f"Adicionando tabela com {len(item.rows)} linhas...")

                num_cols = len(item.headers)
                num_rows = len(item.rows) + 1

                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'

                header_row = table.rows[0]
                for idx, header_text in enumerate(item.headers):
                    cell = header_row.cells[idx]
                    cell.text = header_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'E0E0E0')
                    cell._tc.get_or_add_tcPr().append(shading)

                for row_idx, row_data in enumerate(item.rows):
                    row = table.rows[row_idx + 1]
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = row.cells[col_idx]
                            cell.text = str(cell_text) if cell_text else ""
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(10)

                doc.add_paragraph().space_after = Pt(12)

        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        filename = payload.filename
        if not filename.endswith('.docx'):
            filename += '.docx'
        filename = re.sub(r'[^a-zA-Z0-9\s\-_.]', '', filename)

        print(f"DOCX gerado: {filename}")

        return Response(
            content=doc_buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    except Exception as e:
        print(f"Erro ao gerar DOCX: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erro ao gerar DOCX: {str(e)}")


# ============================================================================
# ENDPOINTS - REVISAO DE ARTIGOS
# ============================================================================

def _is_image_caption(para) -> bool:
    """
    Detecta se um paragrafo e uma legenda de imagem.
    Legendas sao geradas com: alinhamento centralizado, italico, Pt(10), cor cinza (102,102,102).
    """
    if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return False

    runs = para.runs
    if not runs:
        return False

    for run in runs:
        if not run.italic:
            return False
        if run.font.size and run.font.size != Pt(10):
            return False
        if run.font.color.rgb:
            if run.font.color.rgb != RGBColor(102, 102, 102):
                return False

    return True


def _extrair_texto_para_revisao(docx_path: str, incluir_legendas: bool = False) -> tuple:
    """
    Extrai texto estruturado de um DOCX para analise.

    Args:
        docx_path: Caminho do arquivo DOCX
        incluir_legendas: Se True, inclui legendas de imagem (para agente de imagem)

    Returns:
        (texto_completo, titulo)
    """
    doc = Document(docx_path)
    paragrafos = []
    texto_parts = []
    titulo = ""
    idx = 0

    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue

        if _is_image_caption(para):
            if incluir_legendas:
                paragrafos.append({
                    "indice": idx,
                    "texto": texto,
                    "tipo": "image_caption"
                })
                texto_parts.append(f"[P{idx}|IMAGE_CAPTION] {texto}")
                idx += 1
            continue

        estilo = para.style.name if para.style else "Normal"
        tipo = "paragraph"
        if "heading" in estilo.lower():
            tipo = estilo.lower().replace(" ", "")
            if not titulo and "1" in tipo:
                titulo = texto

        paragrafos.append({
            "indice": idx,
            "texto": texto,
            "tipo": tipo
        })

        texto_parts.append(f"[P{idx}|{tipo.upper()}] {texto}")
        idx += 1

    texto_completo = "\n\n".join(texto_parts)
    return texto_completo, titulo


@app.post("/revisao/extrair-texto")
async def revisao_extrair_texto(payload: ExtrairTextoDocxPayload):
    """
    Extrai texto estruturado de um documento DOCX para analise.

    Retorna paragrafos com indices para referencia nas revisoes.
    """
    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)
        paragrafos = []
        texto_parts = []
        idx = 0

        for para in doc.paragraphs:
            texto = para.text.strip()
            if not texto:
                continue

            estilo = para.style.name if para.style else "Normal"
            tipo = "paragraph"
            if "heading" in estilo.lower():
                tipo = estilo.lower().replace(" ", "")

            paragrafos.append({
                "indice": idx,
                "texto": texto,
                "tipo": tipo,
                "tamanho": len(texto)
            })

            texto_parts.append(f"[P{idx}|{tipo.upper()}] {texto}")
            idx += 1

        return {
            "paragrafos": paragrafos,
            "texto_completo": "\n\n".join(texto_parts),
            "total_paragrafos": idx
        }
    finally:
        os.unlink(tmp_path)


@app.post("/revisao/aplicar")
async def revisao_aplicar(payload: AplicarRevisoesPayload):
    """
    Aplica revisoes a um documento DOCX com Track Changes.

    Recebe:
    - docx_url: URL do documento original (ou docx_base64)
    - docx_base64: Documento em base64 (ou docx_url)
    - revisoes: Lista de revisoes no formato:
        {
            "tipo": "SEO|TECNICO|TEXTO",
            "acao": "substituir|deletar|inserir|comentario",
            "texto_original": "texto exato a encontrar",
            "texto_novo": "texto substituto",
            "justificativa": "explicacao"
        }
    - autor: Nome do autor das revisoes

    Retorna: Documento DOCX com Track Changes aplicados
    """
    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        input_path = tmp.name

    output_path = input_path.replace(".docx", "_revisado.docx")

    try:
        revisoes_list = [r.model_dump() for r in payload.revisoes]

        resultado = aplicar_revisoes_docx(
            input_path,
            output_path,
            revisoes_list,
            payload.autor
        )

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="documento_revisado.docx",
            headers={
                "X-Total-Revisoes": str(resultado["total_revisoes"]),
                "X-Aplicadas": str(resultado["aplicadas"]),
                "X-Falhas": str(resultado["falhas"]),
                "X-Comentarios": str(resultado["comentarios"])
            }
        )
    except Exception as e:
        raise HTTPException(500, f"Erro ao aplicar revisoes: {str(e)}")
    finally:
        if os.path.exists(input_path):
            os.unlink(input_path)


@app.post("/revisao/aplicar-json")
async def revisao_aplicar_json(
    docx_url: str = Form(...),
    revisoes: str = Form(...),
    autor: str = Form("Agente IA Revisor")
):
    """
    Aplica revisoes via Form (compativel com n8n HTTP Request).

    Campos Form:
    - docx_url: URL do documento DOCX
    - revisoes: JSON string com array de revisoes
    - autor: Nome do autor (opcional)
    """
    async with httpx.AsyncClient(timeout=60.0) as client:
        resp = await client.get(docx_url)
        if resp.status_code != 200:
            raise HTTPException(400, f"Erro ao baixar documento: {resp.status_code}")

    try:
        revisoes_list = json.loads(revisoes)
    except json.JSONDecodeError as e:
        raise HTTPException(400, f"JSON invalido: {e}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(resp.content)
        input_path = tmp.name

    output_path = input_path.replace(".docx", "_revisado.docx")

    try:
        resultado = aplicar_revisoes_docx(
            input_path,
            output_path,
            revisoes_list,
            autor
        )

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="documento_revisado.docx",
            headers={
                "X-Total-Revisoes": str(resultado["total_revisoes"]),
                "X-Aplicadas": str(resultado["aplicadas"]),
                "X-Falhas": str(resultado["falhas"]),
                "X-Comentarios": str(resultado["comentarios"])
            }
        )
    except Exception as e:
        raise HTTPException(500, f"Erro ao aplicar revisoes: {str(e)}")
    finally:
        if os.path.exists(input_path):
            os.unlink(input_path)


@app.post("/revisao/agente-seo")
async def revisao_agente_seo(payload: RevisaoAgentPayload):
    """
    Executa o agente de revisao SEO.

    Parametros:
    - docx_url: URL do documento DOCX (ou docx_base64)
    - docx_base64: Documento em base64 (ou docx_url)
    - guia_seo_url: URL do guia de SEO (opcional)
    - url_artigo: URL original do artigo (contexto)
    - titulo: Titulo do artigo (contexto)
    """
    print(f"[DEBUG] docx_url: {payload.docx_url[:100] if payload.docx_url else None}...")
    print(f"[DEBUG] docx_base64 len: {len(payload.docx_base64) if payload.docx_base64 else 0}")

    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    guia_seo = "Use boas praticas gerais de SEO para conteudo tecnico educacional."
    if payload.guia_seo_url:
        try:
            async with httpx.AsyncClient(timeout=60.0) as http_client:
                guia_resp = await http_client.get(payload.guia_seo_url)
                if guia_resp.status_code == 200:
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                        tmp.write(guia_resp.content)
                        guia_path = tmp.name
                    guia_doc = Document(guia_path)
                    guia_seo = "\n".join([p.text for p in guia_doc.paragraphs if p.text.strip()])
                    os.unlink(guia_path)
        except Exception as e:
            print(f"Aviso: Nao foi possivel carregar guia SEO: {e}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        if not os.path.exists(tmp_path):
            raise HTTPException(500, f"Arquivo temporario nao foi criado: {tmp_path}")
        if os.path.getsize(tmp_path) == 0:
            raise HTTPException(400, "Documento DOCX vazio ou base64 invalido")

        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo = payload.titulo or titulo_extraido

        system_prompt, user_prompt, artigo_context = formatar_prompt_seo(
            conteudo=conteudo,
            titulo=titulo,
            url=payload.url_artigo or "",
            guia_seo=guia_seo
        )

        llm_client = criar_cliente_llm(provider=payload.provider)
        resposta = llm_client.gerar_resposta(system_prompt, user_prompt, artigo_context=artigo_context)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "SEO"

        return {
            "tipo": "SEO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente SEO: {str(e)}")
    finally:
        os.unlink(tmp_path)


@app.post("/revisao/agente-tecnico")
async def revisao_agente_tecnico(payload: RevisaoAgentPayload):
    """
    Executa o agente de revisao TECNICA.

    Parametros:
    - docx_url: URL do documento DOCX (ou docx_base64)
    - docx_base64: Documento em base64 (ou docx_url)
    - url_artigo: URL original do artigo (contexto)
    - titulo: Titulo do artigo (contexto)
    - data_publicacao: Data de publicacao original
    """
    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        if not os.path.exists(tmp_path):
            raise HTTPException(500, f"Arquivo temporario nao foi criado: {tmp_path}")
        if os.path.getsize(tmp_path) == 0:
            raise HTTPException(400, "Documento DOCX vazio ou base64 invalido")

        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo = payload.titulo or titulo_extraido

        system_prompt, user_prompt, artigo_context = formatar_prompt_tecnico(
            conteudo=conteudo,
            titulo=titulo,
            url=payload.url_artigo or "",
            data_publicacao=payload.data_publicacao or ""
        )

        llm_client = criar_cliente_llm(provider=payload.provider)
        resposta = llm_client.gerar_resposta_com_busca(system_prompt, user_prompt, artigo_context=artigo_context)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "TECNICO"

        return {
            "tipo": "TECNICO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente TECNICO: {str(e)}")
    finally:
        os.unlink(tmp_path)


@app.post("/revisao/agente-texto")
async def revisao_agente_texto(payload: RevisaoAgentPayload):
    """
    Executa o agente de revisao TEXTUAL.

    Parametros:
    - docx_url: URL do documento DOCX (ou docx_base64)
    - docx_base64: Documento em base64 (ou docx_url)
    - url_artigo: URL original do artigo (contexto)
    - titulo: Titulo do artigo (contexto)
    """
    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(400, f"Erro ao obter documento: {str(e)}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        if not os.path.exists(tmp_path):
            raise HTTPException(500, f"Arquivo temporario nao foi criado: {tmp_path}")
        if os.path.getsize(tmp_path) == 0:
            raise HTTPException(400, "Documento DOCX vazio ou base64 invalido")

        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo = payload.titulo or titulo_extraido

        system_prompt, user_prompt, artigo_context = formatar_prompt_texto(
            conteudo=conteudo,
            titulo=titulo,
            url=payload.url_artigo or ""
        )

        llm_client = criar_cliente_llm(provider=payload.provider)
        resposta = llm_client.gerar_resposta(system_prompt, user_prompt, artigo_context=artigo_context)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "TEXTO"

        return {
            "tipo": "TEXTO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente TEXTO: {str(e)}")
    finally:
        os.unlink(tmp_path)


# ============================================================================
# ENDPOINTS - REVISAO VIA FORM (MULTIPART) - Para uso com n8n
# ============================================================================

@app.post("/revisao/agente-seo-form")
async def revisao_agente_seo_form(
    file: UploadFile = File(...),
    provider: str = Form("anthropic"),
    url_artigo: str = Form(""),
    titulo: str = Form(""),
    guia_seo_file: Optional[UploadFile] = File(None),
    palavras_chave: str = Form("")
):
    """
    Executa o agente de revisao SEO via upload de arquivo.
    Ideal para uso com n8n HTTP Request node.

    Parametros:
    - file: Arquivo DOCX do artigo a ser revisado
    - provider: "anthropic" ou "openai"
    - url_artigo: URL original do artigo
    - titulo: Titulo do artigo (opcional, extraido do DOCX se nao fornecido)
    - guia_seo_file: Arquivo DOCX com o guia de SEO da empresa (opcional)
    - palavras_chave: Palavras-chave do Google separadas por virgula ou quebra de linha
    """
    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    guia_path = None

    try:
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo_final = titulo or titulo_extraido

        guia_seo = "Use boas praticas gerais de SEO para conteudo tecnico educacional."
        if guia_seo_file and guia_seo_file.filename:
            guia_bytes = await guia_seo_file.read()
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_guia:
                tmp_guia.write(guia_bytes)
                tmp_guia.flush()
                guia_path = tmp_guia.name

            guia_doc = Document(guia_path)
            guia_seo = "\n".join([p.text for p in guia_doc.paragraphs if p.text.strip()])

        palavras_formatadas = "Nenhuma palavra-chave especifica fornecida. Use seu conhecimento de SEO."
        if palavras_chave and palavras_chave.strip():
            keywords = [kw.strip() for kw in palavras_chave.replace('\n', ',').split(',') if kw.strip()]
            if keywords:
                palavras_formatadas = "\n".join([f"- {kw}" for kw in keywords])

        system_prompt, user_prompt, artigo_context = formatar_prompt_seo(
            conteudo=conteudo,
            titulo=titulo_final,
            url=url_artigo,
            guia_seo=guia_seo,
            palavras_chave=palavras_formatadas
        )

        llm_client = criar_cliente_llm(provider=provider)
        resposta = llm_client.gerar_resposta(system_prompt, user_prompt, artigo_context=artigo_context)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "SEO"

        return {
            "tipo": "SEO",
            "total_sugestoes": len(revisoes),
            "palavras_chave_usadas": palavras_chave.strip() if palavras_chave else None,
            "guia_seo_fornecido": bool(guia_seo_file and guia_seo_file.filename),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente SEO: {str(e)}")
    finally:
        os.unlink(tmp_path)
        if guia_path and os.path.exists(guia_path):
            os.unlink(guia_path)


@app.post("/revisao/agente-tecnico-form")
async def revisao_agente_tecnico_form(
    file: UploadFile = File(...),
    provider: str = Form("anthropic"),
    url_artigo: str = Form(""),
    titulo: str = Form(""),
    data_publicacao: str = Form("")
):
    """
    Executa o agente de revisao TECNICA via upload de arquivo.
    Ideal para uso com n8n HTTP Request node.
    """
    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo_final = titulo or titulo_extraido

        system_prompt, user_prompt, artigo_context = formatar_prompt_tecnico(
            conteudo=conteudo,
            titulo=titulo_final,
            url=url_artigo,
            data_publicacao=data_publicacao
        )

        llm_client = criar_cliente_llm(provider=provider)
        resposta = llm_client.gerar_resposta_com_busca(system_prompt, user_prompt, artigo_context=artigo_context)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "TECNICO"

        return {
            "tipo": "TECNICO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente TECNICO: {str(e)}")
    finally:
        os.unlink(tmp_path)


@app.post("/revisao/agente-texto-form")
async def revisao_agente_texto_form(
    file: UploadFile = File(...),
    provider: str = Form("anthropic"),
    url_artigo: str = Form(""),
    titulo: str = Form("")
):
    """
    Executa o agente de revisao TEXTUAL via upload de arquivo.
    Ideal para uso com n8n HTTP Request node.
    """
    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo_final = titulo or titulo_extraido

        system_prompt, user_prompt, artigo_context = formatar_prompt_texto(
            conteudo=conteudo,
            titulo=titulo_final,
            url=url_artigo
        )

        llm_client = criar_cliente_llm(provider=provider)
        resposta = llm_client.gerar_resposta(system_prompt, user_prompt, artigo_context=artigo_context)
        revisoes = llm_client.extrair_json(resposta)

        for rev in revisoes:
            rev["tipo"] = "TEXTO"

        return {
            "tipo": "TEXTO",
            "total_sugestoes": len(revisoes),
            "revisoes": revisoes
        }

    except Exception as e:
        raise HTTPException(500, f"Erro no agente TEXTO: {str(e)}")
    finally:
        os.unlink(tmp_path)


@app.post("/revisao/agente-imagem")
async def revisao_agente_imagem(payload: RevisaoImagemPayload):
    """
    Executa o agente de revisao de IMAGENS.

    Analisa as imagens do artigo quanto a relevancia, qualidade,
    atualizacao de screenshots e acessibilidade (alt text).

    Parametros:
    - docx_url: URL do documento DOCX (ou docx_base64)
    - docx_base64: Documento em base64 (ou docx_url)
    - url_artigo: URL original do artigo (OBRIGATORIO para extrair imagens)
    - provider: "anthropic" ou "openai" (default: anthropic)
    - titulo: Titulo do artigo (opcional)
    """
    from datetime import datetime

    if not payload.url_artigo:
        raise HTTPException(400, "url_artigo e obrigatorio para o agente de imagem")

    try:
        docx_bytes = await obter_docx_bytes(payload.docx_url, payload.docx_base64)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(docx_bytes)
            tmp.flush()
            tmp_path = tmp.name

        try:
            conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
            titulo_final = payload.titulo or titulo_extraido

            print(f"Extraindo imagens de: {payload.url_artigo}")
            with httpx.Client(timeout=30, follow_redirects=True) as http_client:
                response = http_client.get(payload.url_artigo)
                response.raise_for_status()
                html = response.text

            article_data = extract_article_content(html, payload.url_artigo)

            imagens = [item for item in article_data.get('content', []) if item.get('type') == 'image']
            print(f"{len(imagens)} imagens encontradas")

            if not imagens:
                return {
                    "tipo": "IMAGEM",
                    "total_sugestoes": 0,
                    "revisoes": [],
                    "mensagem": "Nenhuma imagem encontrada no artigo"
                }

            data_atual = datetime.now().strftime("%d/%m/%Y")
            system_prompt, user_prompt, artigo_context = formatar_prompt_imagem(
                conteudo=conteudo,
                imagens=imagens,
                titulo=titulo_final,
                url=payload.url_artigo,
                data_atual=data_atual
            )

            llm_client = criar_cliente_llm(provider=payload.provider)
            print(f"Chamando LLM ({payload.provider}) com {len(imagens)} imagens...")

            try:
                resposta = llm_client.gerar_resposta_com_imagens_e_busca(
                    system_prompt=system_prompt,
                    user_prompt=user_prompt,
                    imagens=imagens,
                    artigo_context=artigo_context
                )
            except Exception as llm_err:
                print(f"Erro ao chamar LLM: {type(llm_err).__name__}: {llm_err}")
                raise

            print(f"Resposta recebida ({len(resposta) if resposta else 0} chars)")
            print(f"Preview: {resposta[:500] if resposta else 'VAZIA'}...")

            revisoes = llm_client.extrair_json(resposta)
            print(f"{len(revisoes)} revisoes extraidas")

            revisoes_validas = []
            for i, rev in enumerate(revisoes):
                if isinstance(rev, dict):
                    rev["tipo"] = "IMAGEM"
                    revisoes_validas.append(rev)
                else:
                    print(f"Revisao ignorada (nao e dict): {type(rev)} - {str(rev)[:100]}")

            return {
                "tipo": "IMAGEM",
                "total_sugestoes": len(revisoes_validas),
                "total_imagens": len(imagens),
                "revisoes": revisoes_validas
            }

        finally:
            os.unlink(tmp_path)

    except httpx.HTTPStatusError as e:
        status = e.response.status_code
        if status >= 500:
            raise HTTPException(502, f"Erro no servidor do artigo ({status}): {e.request.url}")
        else:
            raise HTTPException(400, f"Erro ao buscar URL do artigo ({status}): {e.request.url}")
    except httpx.HTTPError as e:
        raise HTTPException(502, f"Erro de conexao ao buscar artigo: {str(e)}")
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Erro no agente IMAGEM: {str(e)}")


@app.post("/revisao/agente-imagem-form")
async def revisao_agente_imagem_form(
    file: UploadFile = File(...),
    url_artigo: str = Form(...),
    provider: str = Form("anthropic"),
    titulo: str = Form("")
):
    """
    Executa o agente de revisao de IMAGENS via upload de arquivo.
    Ideal para uso com n8n HTTP Request node.

    Parametros:
    - file: Arquivo DOCX
    - url_artigo: URL original do artigo (OBRIGATORIO para extrair imagens)
    - provider: "anthropic" ou "openai" (default: anthropic)
    - titulo: Titulo do artigo (opcional)
    """
    from datetime import datetime

    if not url_artigo:
        raise HTTPException(400, "url_artigo e obrigatorio para o agente de imagem")

    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        conteudo, titulo_extraido = _extrair_texto_para_revisao(tmp_path)
        titulo_final = titulo or titulo_extraido

        print(f"Extraindo imagens de: {url_artigo}")
        with httpx.Client(timeout=30, follow_redirects=True) as http_client:
            response = http_client.get(url_artigo)
            response.raise_for_status()
            html = response.text

        article_data = extract_article_content(html, url_artigo)

        imagens = [item for item in article_data.get('content', []) if item.get('type') == 'image']
        print(f"{len(imagens)} imagens encontradas")

        if not imagens:
            return {
                "tipo": "IMAGEM",
                "total_sugestoes": 0,
                "revisoes": [],
                "mensagem": "Nenhuma imagem encontrada no artigo"
            }

        data_atual = datetime.now().strftime("%d/%m/%Y")
        system_prompt, user_prompt, artigo_context = formatar_prompt_imagem(
            conteudo=conteudo,
            imagens=imagens,
            titulo=titulo_final,
            url=url_artigo,
            data_atual=data_atual
        )

        llm_client = criar_cliente_llm(provider=provider)
        print(f"Chamando LLM ({provider}) com {len(imagens)} imagens...")

        try:
            resposta = llm_client.gerar_resposta_com_imagens_e_busca(
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                imagens=imagens,
                artigo_context=artigo_context
            )
        except Exception as llm_err:
            print(f"Erro ao chamar LLM: {type(llm_err).__name__}: {llm_err}")
            raise

        print(f"Resposta recebida ({len(resposta) if resposta else 0} chars)")

        revisoes = llm_client.extrair_json(resposta)
        print(f"{len(revisoes)} revisoes extraidas")

        revisoes_validas = []
        for rev in revisoes:
            if isinstance(rev, dict):
                rev["tipo"] = "IMAGEM"
                revisoes_validas.append(rev)
            else:
                print(f"Revisao ignorada (nao e dict): {type(rev)} - {str(rev)[:100]}")

        return {
            "tipo": "IMAGEM",
            "total_sugestoes": len(revisoes_validas),
            "total_imagens": len(imagens),
            "revisoes": revisoes_validas
        }

    except httpx.HTTPStatusError as e:
        status = e.response.status_code
        if status >= 500:
            raise HTTPException(502, f"Erro no servidor do artigo ({status}): {e.request.url}")
        else:
            raise HTTPException(400, f"Erro ao buscar URL do artigo ({status}): {e.request.url}")
    except httpx.HTTPError as e:
        raise HTTPException(502, f"Erro de conexao ao buscar artigo: {str(e)}")
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Erro no agente IMAGEM: {str(e)}")
    finally:
        os.unlink(tmp_path)


@app.post("/revisao/aplicar-form")
async def revisao_aplicar_form(
    file: UploadFile = File(...),
    revisoes: str = Form(...),
    autor: str = Form("Agente IA Revisor")
):
    """
    Aplica revisoes a um documento DOCX via upload de arquivo.
    Ideal para uso com n8n HTTP Request node.

    - file: Arquivo DOCX
    - revisoes: JSON string com lista de revisoes
    - autor: Nome do autor das revisoes
    """
    import json as json_lib

    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        input_path = tmp.name

    output_path = input_path.replace(".docx", "_revisado.docx")

    try:
        revisoes_list = json_lib.loads(revisoes)

        resultado = aplicar_revisoes_docx(
            input_path,
            output_path,
            revisoes_list,
            autor
        )

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="documento_revisado.docx",
            headers={
                "X-Total-Revisoes": str(resultado["total_revisoes"]),
                "X-Aplicadas": str(resultado["aplicadas"]),
                "X-Falhas": str(resultado["falhas"]),
                "X-Comentarios": str(resultado["comentarios"])
            }
        )
    except json_lib.JSONDecodeError as e:
        raise HTTPException(400, f"JSON de revisoes invalido: {str(e)}")
    except Exception as e:
        raise HTTPException(500, f"Erro ao aplicar revisoes: {str(e)}")
    finally:
        if os.path.exists(input_path):
            os.unlink(input_path)


@app.post("/revisao/aplicar-comentarios-form")
async def revisao_aplicar_comentarios_form(
    file: UploadFile = File(...),
    revisoes: str = Form(...),
    autor: str = Form("Agente IA Revisor")
):
    """
    Aplica SOMENTE comentarios a um documento DOCX via upload de arquivo.
    Nao altera o texto do documento - apenas adiciona comentarios nos trechos encontrados.
    Suporta multiplos comentarios no mesmo trecho (ranges sobrepostos).

    - file: Arquivo DOCX
    - revisoes: JSON string com lista de revisoes
    - autor: Nome do autor dos comentarios
    """
    import json as json_lib

    original_filename = file.filename or "documento.docx"

    docx_bytes = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        input_path = tmp.name

    output_path = input_path.replace(".docx", "_comentado.docx")

    try:
        revisoes_list = json_lib.loads(revisoes)

        resultado = aplicar_comentarios_docx(
            input_path,
            output_path,
            revisoes_list,
            autor
        )

        stats = resultado.get('estatisticas', {})

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=original_filename,
            headers={
                "X-Total-Comentarios": str(resultado.get('total_comentarios', 0)),
                "X-Match-Exato": str(stats.get('exato', 0)),
                "X-Match-Normalizado": str(stats.get('normalizado', 0)),
                "X-Match-Fuzzy": str(stats.get('fuzzy', 0)),
                "X-Match-Paragrafo": str(stats.get('paragrafo', 0)),
                "X-Match-Falhas": str(stats.get('falha', 0)),
            }
        )
    except json_lib.JSONDecodeError as e:
        raise HTTPException(400, f"JSON de revisoes invalido: {str(e)}")
    except Exception as e:
        raise HTTPException(500, f"Erro ao aplicar comentarios: {str(e)}")
    finally:
        if os.path.exists(input_path):
            os.unlink(input_path)
